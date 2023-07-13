from openpyxl import load_workbook
from docx import Document
import docx
from docx.oxml.ns import qn
from docx2pdf import convert

excel_path = r"C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\수료명단.xlsx"

wb = load_workbook(excel_path, data_only=True)
ws = wb.active

name_list = []
birthday_list = []
ho_list = []
for i in range(2, ws.max_row + 1):
    name_list.append(ws.cell(i, 1).value)
    birthday_list.append(ws.cell(i, 2).value)
    ho_list.append(ws.cell(i, 3).value)

print("이름: ", name_list)
print("생일:", birthday_list)
print("호:", ho_list)

doc = Document(r'C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\수료증양식.docx')

for i in range(len(name_list)):
    doc.paragraphs[3].clear()
    run = doc.paragraphs[3].add_run('제' + ho_list[i] + '호')
    run.font.name = '나눔고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    doc.paragraphs[6].clear()
    run = doc.paragraphs[6].add_run('성        명:' + name_list[i])
    run.font.name = '나눔고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(18)

    doc.paragraphs[7].clear()
    run = doc.paragraphs[7].add_run('생 년 월 일:' + birthday_list[i])
    run.font.name = '나눔고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(18)

    doc.save(r"C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\\" + ho_list[i] + name_list[i] + ".docx")
    convert(r"C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\\" + ho_list[i] + name_list[i] + ".docx", r"C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\\" + ho_list[i] + name_list[i] + ".pdf")
