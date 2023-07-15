from docx import Document
import docx
from docx.oxml.ns import qn

doc = Document(r'C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\수료증양식.docx')

for i, paragraph in enumerate(doc.paragraphs):
    print(str(i) + ":" + paragraph.text)

doc.paragraphs[3].clear()
run = doc.paragraphs[3].add_run('제 2022-9999 호')
run.font.name = '나눔고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

doc.paragraphs[6].clear()
run = doc.paragraphs[6].add_run('성        명: 장문철')
run.font.name = '나눔고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(18)

doc.paragraphs[7].clear()
run = doc.paragraphs[7].add_run('생 년 월 일: 1984.09.19')
run.font.name = '나눔고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(18)

doc.save(r'C:\Users\USER\PycharmProjects\pythonProject3\전산 프로젝트\수료증 자동생성 후 PDF 변환\수료증수정.docx')
